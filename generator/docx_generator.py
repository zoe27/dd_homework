"""
Word 文档生成器

将 HomeworkDocument 生成 .docx 文件，支持文字排版 + 图片嵌入。

可直接运行测试：
  python -m generator.docx_generator
"""

import os
import sys
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.models import HomeworkDocument, HomeworkCard
from utils.logger import logger
import config

# 图片最大宽度（cm）
IMAGE_MAX_WIDTH_CM = 14


def _add_title(doc: Document, homework_date: date) -> None:
    """添加标题和日期副标题"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("今日作业")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    weekdays = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    weekday = weekdays[homework_date.weekday()]
    date_str = f"{homework_date.year}年{homework_date.month}月{homework_date.day}日（{weekday}）"
    run = subtitle.add_run(date_str)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 空行


def _add_card(doc: Document, card: HomeworkCard) -> None:
    """添加单科作业块：科目标题 + 作业列表 + 图片"""
    # 科目标题
    subject_para = doc.add_paragraph()
    run = subject_para.add_run(f"【{card.subject}】")
    run.bold = True
    run.font.size = Pt(14)

    # 作业条目
    for i, item in enumerate(card.items, 1):
        item_para = doc.add_paragraph()
        item_para.paragraph_format.left_indent = Cm(0.5)
        item_para.add_run(f"{i}. {item}").font.size = Pt(12)

    # 图片
    for img_path in card.image_paths:
        if not os.path.exists(img_path):
            logger.warning(f"图片不存在，跳过: {img_path}")
            continue
        try:
            doc.add_picture(img_path, width=Cm(IMAGE_MAX_WIDTH_CM))
            logger.info(f"已插入图片: {img_path}")
        except Exception as e:
            logger.warning(f"图片插入失败: {img_path} — {e}")

    doc.add_paragraph()  # 科目间空行


def generate(document: HomeworkDocument) -> str:
    """
    生成 Word 文档，返回文件路径。
    """
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # 页面边距（适合 A4 打印）
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_title(doc, document.date)

    for card in document.cards:
        _add_card(doc, card)

    date_str = document.date.strftime("%Y%m%d")
    filename = f"今日作业_{date_str}.docx"
    filepath = os.path.join(config.OUTPUT_DIR, filename)
    doc.save(filepath)
    logger.info(f"文档已保存: {filepath}")
    return filepath


# ── CLI 测试入口 ──────────────────────────────────────────────────────────────

def main():
    from models.models import HomeworkCard, HomeworkDocument

    today = date.today()
    cards = [
        HomeworkCard(date=today, subject="数学",
                     items=["订正知能P17,P18", "完成知能P19,P20"]),
        HomeworkCard(date=today, subject="语文",
                     items=["背诵第三课课文", "完成练习册第5页"]),
        HomeworkCard(date=today, subject="英语",
                     items=["听写单词20个", "朗读课文三遍"]),
    ]
    doc = HomeworkDocument(date=today, cards=cards)
    path = generate(doc)
    print(f"\n文档已生成: {path}")


if __name__ == "__main__":
    main()
